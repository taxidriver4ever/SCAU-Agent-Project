import os
import base64
import io
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import Inches
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from PIL import Image
import numpy as np
from openai import OpenAI
import json
import time
from pathlib import Path

class ImageExtractor:
    """从Word文档中提取图片及其上下文的类"""
    
    def __init__(self, debug_folder: str):
        self.debug_folder = debug_folder# 初始化Qwen2.5-VL客户端
        self.qwen_vl_client = OpenAI(
            base_url='https://api-inference.modelscope.cn/v1',
            api_key='ms-87b5cf21-1038-42ef-8f98-307a1eb45cc8'
        )
        
        # 初始化Qwen3.0客户端
        self.qwen3_client = OpenAI(
            base_url='https://api-inference.modelscope.cn/v1',
            api_key='ms-87b5cf21-1038-42ef-8f98-307a1eb45cc8'
        )
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """从Word文档中提取图片及其上下文"""
        doc = Document(docx_path)
        images_data = []
        
        # 获取所有段落文本
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 遍历文档中的所有段落，查找包含图片的段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                # 查找图片元素
                for drawing in run.element.xpath('.//a:blip'):
                    try:
                        # 获取图片的关系ID
                        embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id:
                            # 通过关系ID获取图片数据
                            image_part = doc.part.related_parts[embed_id]
                            image_data = image_part.blob
                            
                            # 获取上下文
                            context_before = ""
                            context_after = ""
                            
                            # 获取图片前后的段落作为上下文
                            if para_idx > 0:
                                context_before = " ".join([p.text.strip() for p in doc.paragraphs[max(0, para_idx-2):para_idx] if p.text.strip()])[:200]
                            if para_idx < len(doc.paragraphs) - 1:
                                context_after = " ".join([p.text.strip() for p in doc.paragraphs[para_idx+1:min(len(doc.paragraphs), para_idx+3)] if p.text.strip()])[:200]
                            
                            images_data.append({
                                'image_data': image_data,
                                'context_before': context_before,
                                'context_after': context_after,
                                'source_file': os.path.basename(docx_path)
                            })
                    except Exception as e:
                        print(f"提取图片时出错: {e}")
                        continue
        
        return images_data
    
    def image_to_base64(self, image_data: bytes) -> str:
        """将图片数据转换为base64编码"""
        return base64.b64encode(image_data).decode('utf-8')
    
    def describe_image_with_qwen_vl(self, image_data: bytes) -> str:
        """使用Qwen2.5-VL识别图片内容"""
        try:
            # 将图片转换为base64
            base64_image = self.image_to_base64(image_data)
            image_url = f"data:image/jpeg;base64,{base64_image}"
            
            response = self.qwen_vl_client.chat.completions.create(
                model='Qwen/Qwen2.5-VL-72B-Instruct',
                messages=[{
                    'role': 'user',
                    'content': [{
                        'type': 'text',
                        'text': '请详细描述这幅图片的内容，包括图片中的文字、图形、布局等所有可见元素。',
                    }, {
                        'type': 'image_url',
                        'image_url': {
                            'url': image_url,
                        },
                    }],
                }],
                stream=False
            )
            
            return response.choices[0].message.content
        except Exception as e:
            print(f"Qwen2.5-VL识别图片时出错: {e}")
            return "图片识别失败"
    
    def enhance_description_with_qwen3(self, image_description: str, context_before: str, context_after: str) -> str:
        """使用Qwen3.0结合上下文完善图片描述"""
        try:
            prompt = f"""
            请根据以下信息，生成一段完整的图片内容描述：
            
            图片识别结果：{image_description}
            
            上文内容：{context_before}
            
            下文内容：{context_after}
            
            请结合上下文，分析这张图片的应用场景、主题和作用，并生成一段完整、准确的描述。描述应该包括：
            1. 图片的具体内容
            2. 图片在文档中的作用和意义
            3. 与上下文的关联性
            4. 应用场景和主题
            
            请用一段话总结，不超过200字。
            """
            
            response = self.qwen3_client.chat.completions.create(
                model='Qwen/Qwen3-235B-A22B-Instruct-2507',
                messages=[
                    {
                        'role': 'system',
                        'content': 'You are a helpful assistant that analyzes images in context.'
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ],
                stream=False
            )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Qwen3.0完善描述时出错: {e}")
            return image_description  # 返回原始描述作为备选
    
    def save_images_to_word(self, processed_images: List[Dict], output_path: str):
        """将处理后的图片和描述保存到Word文档"""
        doc = Document()
        doc.add_heading('提取的图片及其描述', 0)
        
        for i, img_data in enumerate(processed_images):
            # 添加图片标题
            doc.add_heading(f'图片 {i+1} - 来源: {img_data["source_file"]}', level=1)
            
            # 添加图片
            try:
                # 将图片数据保存为临时文件
                temp_image_path = f"temp_image_{i}.png"
                with open(temp_image_path, 'wb') as f:
                    f.write(img_data['image_data'])
                
                # 添加图片到文档
                doc.add_picture(temp_image_path, width=Inches(4))
                
                # 删除临时文件
                os.remove(temp_image_path)
            except Exception as e:
                doc.add_paragraph(f"图片加载失败: {e}")
            
            # 添加上下文
            doc.add_heading('上下文', level=2)
            doc.add_paragraph(f"上文: {img_data['context_before']}")
            doc.add_paragraph(f"下文: {img_data['context_after']}")
            
            # 添加描述
            doc.add_heading('图片描述', level=2)
            doc.add_paragraph(img_data['enhanced_description'])
            
            # 添加分隔线
            doc.add_paragraph('\n' + '='*50 + '\n')
        
        doc.save(output_path)
        print(f"图片和描述已保存到: {output_path}")
    
    def process_all_documents(self) -> List[Dict]:
        """处理debug文件夹中的所有Word文档"""
        all_processed_images = []
        
        # 获取所有docx文件
        docx_files = [f for f in os.listdir(self.debug_folder) if f.endswith('.docx')]
        
        print(f"找到 {len(docx_files)} 个Word文档")
        
        for docx_file in docx_files:
            print(f"\n处理文档: {docx_file}")
            docx_path = os.path.join(self.debug_folder, docx_file)
            
            try:
                # 提取图片
                images_data = self.extract_images_from_docx(docx_path)
                print(f"从 {docx_file} 中提取到 {len(images_data)} 张图片")
                
                for i, img_data in enumerate(images_data):
                    print(f"  处理图片 {i+1}/{len(images_data)}...")
                    
                    # 使用Qwen2.5-VL识别图片
                    image_description = self.describe_image_with_qwen_vl(img_data['image_data'])
                    print(f"    图片识别完成")
                    
                    # 使用Qwen3.0完善描述
                    enhanced_description = self.enhance_description_with_qwen3(
                        image_description, 
                        img_data['context_before'], 
                        img_data['context_after']
                    )
                    print(f"    描述完善完成")
                    
                    # 添加到结果列表
                    processed_img = {
                        'image_data': img_data['image_data'],
                        'context_before': img_data['context_before'],
                        'context_after': img_data['context_after'],
                        'source_file': img_data['source_file'],
                        'original_description': image_description,
                        'enhanced_description': enhanced_description
                    }
                    all_processed_images.append(processed_img)
                    
                    # 添加延迟避免API限制
                    time.sleep(1)
                    
            except Exception as e:
                print(f"处理文档 {docx_file} 时出错: {e}")
                continue
        
        return all_processed_images

def main():
    """主函数"""
    debug_folder = "d:\\core code\\chunkit_fronted\\debug"
    output_word_path = "d:\\core code\\chunkit_fronted\\extracted_images.docx"
    
    # 创建图片提取器
    extractor = ImageExtractor(debug_folder)
    
    print("开始处理Word文档中的图片...")
    
    # 处理所有文档
    processed_images = extractor.process_all_documents()
    
    if processed_images:
        print(f"\n总共处理了 {len(processed_images)} 张图片")
        
        # 保存到Word文档
        extractor.save_images_to_word(processed_images, output_word_path)
        
        # 保存处理结果到JSON文件（用于后续FAISS数据库更新）
        json_output_path = "d:\\core code\\chunkit_fronted\\processed_images.json"
        with open(json_output_path, 'w', encoding='utf-8') as f:
            # 将图片数据转换为base64以便JSON序列化
            json_data = []
            for img in processed_images:
                json_img = img.copy()
                json_img['image_data_base64'] = base64.b64encode(img['image_data']).decode('utf-8')
                del json_img['image_data']  # 删除原始二进制数据
                json_data.append(json_img)
            
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        print(f"处理结果已保存到JSON文件: {json_output_path}")
        print("\n图片提取和处理完成！")
    else:
        print("没有找到任何图片")

if __name__ == "__main__":
    main()